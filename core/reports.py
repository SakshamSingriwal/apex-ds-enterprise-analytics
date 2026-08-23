"""
Report export: PDF (reportlab) and DOCX (python-docx).
"""
from __future__ import annotations

import io
import logging
from typing import Optional, Any

import pandas as pd
import numpy as np

logger = logging.getLogger("apex_ds.reports")


def export_report(
    df: pd.DataFrame,
    target: Any,
    problem_type: Any,
    model: Optional[Any] = None,
    format: str = "pdf",
    include_charts: bool = True,
) -> bytes:
    """Generate report and return as bytes."""
    if format.lower() == "pdf":
        return _build_pdf(df, target, problem_type, model, include_charts)
    else:
        return _build_docx(df, target, problem_type, model, include_charts)


# ── PDF ────────────────────────────────────────────────────────────────────

def _build_pdf(df: pd.DataFrame, target: Any, problem_type: Any, model: Any, include_charts: bool) -> bytes:
    from reportlab.lib.pagesizes import A4  # type: ignore[import]
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore[import]
    from reportlab.lib.units import cm  # type: ignore[import]
    from reportlab.lib import colors  # type: ignore[import]
    from reportlab.platypus import (  # type: ignore[import]
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=20, textColor=colors.HexColor("#1f6feb"))
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#388bfd"))
    body = styles["BodyText"]

    story = []

    # Header
    story.append(Paragraph("Apex DS – Enterprise Analytics Report", title_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#30363d")))
    story.append(Spacer(1, 0.5 * cm))

    # Dataset info
    story.append(Paragraph("1. Dataset Overview", h2_style))
    rows, cols = df.shape
    missing_pct = round(float(df.isnull().sum().sum() / (rows * cols) * 100), 2)
    data_table = [
        ["Metric", "Value"],
        ["Rows", f"{rows:,}"],
        ["Columns", str(cols)],
        ["Target column", str(target)],
        ["Problem type", str(problem_type).replace("_", " ").title()],
        ["Missing %", f"{missing_pct:.2f}%"],
    ]
    tbl = Table(data_table, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f6feb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f6f8fa"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7de")),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.5 * cm))

    # Model section
    story.append(Paragraph("2. Model Results", h2_style))
    if model is not None:
        try:
            lb = model.leaderboard(silent=True)
            best = lb.iloc[0]
            story.append(Paragraph(f"Best model: <b>{best['model']}</b>", body))
            story.append(Paragraph(f"Validation score: <b>{best['score_val']:.4f}</b>", body))
            story.append(Spacer(1, 0.3 * cm))
        except Exception:
            story.append(Paragraph("Model leaderboard unavailable.", body))

        # Feature importance
        if include_charts:
            try:
                fi = model.feature_importance(df)
                if fi is not None:
                    col = "importance" if "importance" in fi.columns else fi.columns[0]
                    top10 = fi.head(10)
                    fi_data = [["Feature", "Importance"]]
                    for feat, imp in zip(top10.index, top10[col]):
                        fi_data.append([str(feat), f"{float(imp):.4f}"])
                    fi_tbl = Table(fi_data, hAlign="LEFT")
                    fi_tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#388bfd")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f6f8fa"), colors.white]),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7de")),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ]))
                    story.append(Paragraph("Top 10 Feature Importances:", h2_style))
                    story.append(fi_tbl)
                    story.append(Spacer(1, 0.3 * cm))
            except Exception:
                pass
    else:
        story.append(Paragraph("No trained model found. Train a model first via the AutoML tab.", body))

    story.append(Spacer(1, 0.5 * cm))

    # Business insights
    story.append(Paragraph("3. Business Insights & Recommendations", h2_style))
    from core.business_insights import generate_business_insights
    insights = generate_business_insights(df, target, problem_type, model)
    story.append(Paragraph(insights["summary"], body))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("<b>Recommendations:</b>", body))
    for rec in insights["recommendations"]:
        story.append(Paragraph(f"• {rec}", body))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("<b>Risk Factors:</b>", body))
    for risk in insights["risks"]:
        story.append(Paragraph(f"⚠ {risk}", body))

    # Footer
    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#30363d")))
    story.append(Paragraph("Generated by Apex DS – Enterprise Analytics Studio", styles["Italic"]))

    doc.build(story)
    return buf.getvalue()


# ── DOCX ───────────────────────────────────────────────────────────────────

def _build_docx(df: pd.DataFrame, target: Any, problem_type: Any, model: Any, include_charts: bool) -> bytes:
    from docx import Document  # type: ignore[import]
    from docx.shared import Pt, RGBColor  # type: ignore[import]

    doc = Document()
    doc.add_heading("Apex DS – Enterprise Analytics Report", level=0)

    # Dataset info
    doc.add_heading("1. Dataset Overview", level=1)
    rows, cols = df.shape
    missing_pct = round(float(df.isnull().sum().sum() / (rows * cols) * 100), 2)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    data = [
        ("Rows", f"{rows:,}"),
        ("Columns", str(cols)),
        ("Target", str(target)),
        ("Problem Type", str(problem_type).replace("_", " ").title()),
        ("Missing %", f"{missing_pct:.2f}%"),
        ("Numeric Features", str(len(df.select_dtypes(include=[np.number]).columns))),
    ]
    for i, (k, v) in enumerate(data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    # Model section
    doc.add_heading("2. Model Results", level=1)
    if model is not None:
        try:
            lb = model.leaderboard(silent=True)
            best = lb.iloc[0]
            doc.add_paragraph(f"Best model: {best['model']}")
            doc.add_paragraph(f"Validation score: {best['score_val']:.4f}")
        except Exception:
            doc.add_paragraph("Model leaderboard unavailable.")
    else:
        doc.add_paragraph("No trained model found.")

    # Insights
    doc.add_heading("3. Business Insights", level=1)
    from core.business_insights import generate_business_insights
    insights = generate_business_insights(df, target, problem_type, model)
    doc.add_paragraph(insights["summary"])
    doc.add_heading("Recommendations", level=2)
    for rec in insights["recommendations"]:
        doc.add_paragraph(f"• {rec}")
    doc.add_heading("Risk Factors", level=2)
    for risk in insights["risks"]:
        doc.add_paragraph(f"⚠ {risk}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()